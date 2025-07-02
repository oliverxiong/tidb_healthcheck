import requests
import os
import sys
import datetime
from docx import Document
from docx.shared import Cm
import configparser
import logging
import subprocess
import re
import argparse

################################################################################################
# TiDB Grafana Panel Mapping
#overview eDbRZpnWk
#76 Services Port Status
#27 Storage capacity
#28 Current storage size
#30 Number of Regions
#65 Normal stores
#18 Abnormal stores
#66 Region health
#34 Duration
#20 Leader
#21 Region
#68 Hot write Region's leader distribution
#69 Hot read Region's leader distribution
#2 Statement OPS

#55 CPU Usage
#58 Memory Available
#79 Network Traffic
#61 IO Util

#pd Q6RuHYIWk
#83 Store capacity
#91 Store available
#41 Store Region score
#40 Store leader score
#46 Scheduler is running
#45 Schedule operator create
#78 Schedule operator timeout
#77 Schedule operator finish

#TIDB 000000011
#184 Uptime
#8 Connection Count
#111 Get Token Duration
#191 Skip Binlog Count
#156 Parse Duration
#154 Compile Duration
#77 PD TSO Wait Duration
#78 PD TSO RPC Duration
#53 KV Backoff OPS
#11 TiClient Region Error OPS
#32 Lock Resolve OPS


#tikv RDVQiEzZz
#61 Raft store CPU
#79 Async apply CPU
#64 Scheduler worker CPU
#105 gRPC poll CPU
#4287 Unified read pool CPU
#109 Storage async write duration
#3062 Request duration
#2741 Critical error
#1584 Server is busy
#18 Server report failures
#1718 Raftstore error
#1719 Scheduler error

#tiflash SVbh2xUWk
#11 Request Duration
#13 Request Handle Duration

#TiFlash-Proxy-Details  kWxNAVnGz
#1714 Region

#SVbh2xUWk TiFlash-Summary
#12 Error QPS

#YiGL8hBZ1 tidb-test-ticdc
#398 Changefeed checkpoint lag
#468 Changefeed resolved ts lag

################################################################################################
# create parser
parser = argparse.ArgumentParser(description='This tool is using for automation generate TiDB Health Check report。')
# 添加参数
parser.add_argument('-v', '--version', action='version', version='%(prog)s 1.1')
parser.add_argument('-c', '--cfg', type=str, help='set config file path', default ='tidb_healthcheck.cfg')
#parser.add_argument('-t', '--template', type=str, help='set report template file path', default ='TiDB健康检查报告-template')
args = parser.parse_args()
if os.path.exists(args.cfg):
    config = configparser.ConfigParser()
    config.read(args.cfg)
    print(f"Using config file - {args.cfg}")
else:
    print(f"Config file - {args.cfg} is not exists, please provide correct config file path !")
    sys.exit()

logging.basicConfig(filename='tidb_healthcheck.log', level=logging.INFO, format='%(asctime)s %(levelname)s - %(message)s')
# Grafana API address
GRAFANA_URL = config.get('MONITOR', 'GRAFANA_URL')
# Grafana API Keys
GRAFANA_API_KEY = config.get('MONITOR', 'GRAFANA_API_KEY')
# Define Image Dir
OUTPUT_DIR = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
FROM_TIME = config.get('MONITOR', 'FROM_TIME')
#TO_TIME = "now"
TO_TIME = config.get('MONITOR', 'TO_TIME')
# IMG_WIDTH = 1000
IMG_WIDTH = config.get('MONITOR', 'IMG_WIDTH')
#IMG_HEIGHT = 500
IMG_HEIGHT = config.get('MONITOR', 'IMG_HEIGHT')
# ORGID = 1
ORGID = config.get('MONITOR', 'ORGID')
REPORT_TEMPLATE = config.get('DEFAULT', 'REPORT_TEMPLATE')
if not os.path.exists(REPORT_TEMPLATE):
    print(f"Report template file - {REPORT_TEMPLATE} is not exists !")
    sys.exit()
CLUSTER_NAME = config.get('CLUSTER', 'CLUSTER_NAME')
TIDB_SERVER_IP = config.get('DATABASE', 'TIDB_SERVER_IP')
TIDB_SERVER_PORT = config.get('DATABASE', 'TIDB_SERVER_PORT')
TIDB_USER = config.get('DATABASE', 'TIDB_USER')
TIDB_USER_PWD = config.get('DATABASE', 'TIDB_USER_PWD')
OUTPUT_REPORT = REPORT_TEMPLATE.replace(".docx", f'.{CLUSTER_NAME}.{datetime.datetime.now().strftime("%Y-%m-%d")}.docx')

PANELS_TO_RENDER = [
    {"dashboard_id": "eDbRZpnWk", "panel_id": 76, "output_file": "overview.Services_Port_Status.png"},
    {"dashboard_id": "eDbRZpnWk", "panel_id": 27, "output_file": "overview.Storage_capacity.png"},
    {"dashboard_id": "eDbRZpnWk", "panel_id": 28, "output_file": "overview.Current_storage_size.png"},
    {"dashboard_id": "eDbRZpnWk", "panel_id": 30, "output_file": "overview.Number_of_Regions.png"},
    {"dashboard_id": "eDbRZpnWk", "panel_id": 65, "output_file": "overview.Normal_stores.png"},
    {"dashboard_id": "eDbRZpnWk", "panel_id": 18, "output_file": "overview.Abnormal_stores.png"},
    {"dashboard_id": "eDbRZpnWk", "panel_id": 66, "output_file": "overview.Region_health.png"},
    {"dashboard_id": "eDbRZpnWk", "panel_id": 34, "output_file": "overview.Duration.png"},
    {"dashboard_id": "eDbRZpnWk", "panel_id": 20, "output_file": "overview.Leader.png"},
    {"dashboard_id": "eDbRZpnWk", "panel_id": 21, "output_file": "overview.Region.png"},
    {"dashboard_id": "eDbRZpnWk", "panel_id": 68, "output_file": "overview.Hot_write_Regions_leader_distribution.png"},
    {"dashboard_id": "eDbRZpnWk", "panel_id": 69, "output_file": "overview.Hot_read_Regions_leader_distribution.png"},
    {"dashboard_id": "eDbRZpnWk", "panel_id": 2, "output_file": "overview.Statement_OPS.png"},
    {"dashboard_id": "eDbRZpnWk", "panel_id": 55, "output_file": "overview.CPU_Usage.png"},
    {"dashboard_id": "eDbRZpnWk", "panel_id": 58, "output_file": "overview.Memory_Available.png"},
    {"dashboard_id": "eDbRZpnWk", "panel_id": 79, "output_file": "overview.Network_Traffic.png"},
    {"dashboard_id": "eDbRZpnWk", "panel_id": 61, "output_file": "overview.IO_Util.png"},
    {"dashboard_id": "Q6RuHYIWk", "panel_id": 83, "output_file": "PD.Store_capacity.png"},
    {"dashboard_id": "Q6RuHYIWk", "panel_id": 91, "output_file": "PD.Store_available.png"},
    {"dashboard_id": "Q6RuHYIWk", "panel_id": 41, "output_file": "PD.Store_Region_score.png"},
    {"dashboard_id": "Q6RuHYIWk", "panel_id": 40, "output_file": "PD.Store_leader_score.png"},
    {"dashboard_id": "Q6RuHYIWk", "panel_id": 46, "output_file": "PD.Scheduler_is_running.png"},
    {"dashboard_id": "Q6RuHYIWk", "panel_id": 45, "output_file": "PD.Schedule_operator_create.png"},
    {"dashboard_id": "Q6RuHYIWk", "panel_id": 78, "output_file": "PD.Schedule_operator_timeout.png"},
    {"dashboard_id": "Q6RuHYIWk", "panel_id": 77, "output_file": "PD.Schedule_operator_finish.png"},
    {"dashboard_id": "000000011", "panel_id": 184, "output_file": "TIDB.Uptime.png"},
    {"dashboard_id": "000000011", "panel_id": 8, "output_file": "TIDB.Connection_Count.png"},
    {"dashboard_id": "000000011", "panel_id": 111, "output_file": "TIDB.Get_Token_Duration.png"},
    {"dashboard_id": "000000011", "panel_id": 191, "output_file": "TIDB.Skip_Binlog_Count.png"},
    {"dashboard_id": "000000011", "panel_id": 156, "output_file": "TIDB.Parse_Duration.png"},
    {"dashboard_id": "000000011", "panel_id": 154, "output_file": "TIDB.Compile_Duration.png"},
    {"dashboard_id": "000000011", "panel_id": 77, "output_file": "TIDB.PD_TSO_Wait_Duration.png"},
    {"dashboard_id": "000000011", "panel_id": 78, "output_file": "TIDB.PD_TSO_RPC_Duration.png"},
    {"dashboard_id": "000000011", "panel_id": 53, "output_file": "TIDB.KV_Backoff_OPS.png"},
    {"dashboard_id": "000000011", "panel_id": 11, "output_file": "TIDB.TiClient_Region_Error_OPS.png"},
    {"dashboard_id": "000000011", "panel_id": 32, "output_file": "TIDB.Lock_Resolve_OPS.png"},
    {"dashboard_id": "RDVQiEzZz", "panel_id": 61, "output_file": "TIKV-Details.Raft_store_CPU.png"},
    {"dashboard_id": "RDVQiEzZz", "panel_id": 79, "output_file": "TIKV-Details.Async_apply_CPU.png"},
    {"dashboard_id": "RDVQiEzZz", "panel_id": 64, "output_file": "TIKV-Details.Scheduler_worker_CPU.png"},
    {"dashboard_id": "RDVQiEzZz", "panel_id": 105, "output_file": "TIKV-Details.gRPC_poll_CPU.png"},
    {"dashboard_id": "RDVQiEzZz", "panel_id": 4287, "output_file": "TIKV-Details.Unified_read_pool_CPU.png"},
    {"dashboard_id": "RDVQiEzZz", "panel_id": 109, "output_file": "TIKV-Details.Storage_async_write_duration.png"},
    {"dashboard_id": "RDVQiEzZz", "panel_id": 3062, "output_file": "TIKV-Details.Request_duration.png"},
    {"dashboard_id": "RDVQiEzZz", "panel_id": 2741, "output_file": "TIKV-Details.Critical_error.png"},
    {"dashboard_id": "RDVQiEzZz", "panel_id": 1584, "output_file": "TIKV-Details.Server_is_busy.png"},
    {"dashboard_id": "RDVQiEzZz", "panel_id": 18, "output_file": "TIKV-Details.Server_report_failures.png"},
    {"dashboard_id": "RDVQiEzZz", "panel_id": 1718, "output_file": "TIKV-Details.Raftstore_error.png"},
    {"dashboard_id": "RDVQiEzZz", "panel_id": 1719, "output_file": "TIKV-Details.Scheduler_error.png"},
    {"dashboard_id": "SVbh2xUWk", "panel_id": 11, "output_file": "TiFlash.Request_Duration.png"},
    {"dashboard_id": "SVbh2xUWk", "panel_id": 13, "output_file": "TiFlash.Request_Handle_Duration.png"},
    {"dashboard_id": "kWxNAVnGz", "panel_id": 1714, "output_file": "TiFlash-Proxy-Details.Region.png"},
    {"dashboard_id": "SVbh2xUWk", "panel_id": 12, "output_file": "TiFlash-Summary.Error_QPS.png"},
    {"dashboard_id": "YiGL8hBZ1", "panel_id": 398, "output_file": "Changefeed_checkpoint_lag.png"},
    {"dashboard_id": "YiGL8hBZ1", "panel_id": 468, "output_file": "Changefeed_resolved_ts_lag.png"}
]

COMMANDS = [
    {"cmd_name": "tiup_cluster_list", "cmd_text": f'''tiup cluster list'''},
    {"cmd_name": "tiup_cluster_check", "cmd_text": f'''tiup cluster check --cluster {CLUSTER_NAME}'''},
    {"cmd_name": "tiup_cluster_show_config", "cmd_text": f'''tiup cluster show-config {CLUSTER_NAME}'''},
    {"cmd_name": "tiup_cluster_display", "cmd_text": f'''tiup cluster display {CLUSTER_NAME}'''},
    {"cmd_name": "tiup_cluster_df", "cmd_text": f'''tiup cluster exec {CLUSTER_NAME} --command "sudo df -Th"'''},
    #{"cmd_name": "tiup_cluster_eth_speed", "cmd_text": f'''tiup cluster exec {CLUSTER_NAME} --command "ethtool bond0 |grep Speed"'''},
    {"cmd_name": "tiup_cluster_eth_speed", "cmd_text": f'''tiup cluster exec {CLUSTER_NAME} --command "sudo ip link show | grep -Eo "^[0-9]+:.*:" | cut -d: -f2- | xargs -n1 ethtool || echo ..."'''},
    #{"cmd_name": "tiup_cluster_os_ver", "cmd_text": f'''tiup cluster exec {CLUSTER_NAME} --command "sudo nkvers"'''},
    #{"cmd_name": "tiup_cluster_os_ver", "cmd_text": f'''tiup cluster exec {CLUSTER_NAME} --command "sudo lsb_release -a"'''},
    {"cmd_name": "tiup_cluster_os_ver", "cmd_text": f'''tiup cluster exec {CLUSTER_NAME} --command "sudo cat /etc/os-release"'''},
    {"cmd_name": "tiup_cluster_kernel_param", "cmd_text": f'''tiup cluster exec {CLUSTER_NAME} --command "sudo cat /etc/sysctl.conf |grep -v '#'"'''},
    {"cmd_name": "tiup_cluster_ulimit", "cmd_text": f'''tiup cluster exec {CLUSTER_NAME} --command "sudo cat /etc/security/limits.conf  |grep -v '#'"'''},
    {"cmd_name": "tiup_cluster_swap", "cmd_text": f'''tiup cluster exec {CLUSTER_NAME} --command "sudo free -m"'''},
    {"cmd_name": "tiup_cluster_th", "cmd_text": f'''tiup cluster exec {CLUSTER_NAME} --command "sudo cat /sys/kernel/mm/transparent_hugepage/enabled"'''},
    {"cmd_name": "tiup_cluster_ntp", "cmd_text": f'''tiup cluster exec {CLUSTER_NAME} --command "sudo systemctl status chronyd.service"'''},
    {"cmd_name": "tiup_cluster_ext4", "cmd_text": f'''tiup cluster exec {CLUSTER_NAME} --command "sudo mount -t ext4"'''},
    {"cmd_name": "tiup_cluster_firewall", "cmd_text": f'''tiup cluster exec {CLUSTER_NAME} --command "sudo systemctl status firewalld;echo"'''},
    {"cmd_name": "tiup_cluster_cpupower", "cmd_text": f'''tiup cluster exec {CLUSTER_NAME} --command "sudo cpupower frequency-info | grep policy"'''},
    {"cmd_name": "tiup_cluster_netbonding", "cmd_text": f'''tiup cluster exec {CLUSTER_NAME} --command "sudo cat /proc/net/bonding/bond0|grep Bonding || echo ..."'''},
    {"cmd_name": "tiup_cluster_display_pump_drainer", "cmd_text": f'''tiup cluster display {CLUSTER_NAME} -R pump,drainer'''},
    {"cmd_name": "tiup_cluster_display_cdc", "cmd_text": f'''tiup cluster display {CLUSTER_NAME} -R cdc'''},
    #{"cmd_name": "tiup_cluster_IO_scheduler", "cmd_text": f'''tiup cluster exec {CLUSTER_NAME} --command "sudo cat /sys/block/*/queue/scheduler"'''},
    {"cmd_name": "tiup_cluster_IO_scheduler", "cmd_text": f'''tiup cluster exec {CLUSTER_NAME} --command "for scheduler in /sys/block/*/queue/scheduler; do echo "$(dirname "$scheduler")"; cat "$scheduler"; done"'''},
    #{"cmd_name": "tiup_cdc_cf_list", "cmd_text": f'''tiup cdc cli changefeed list'''},
    #{"cmd_name": "mysql.user", "cmd_text": f'mysql -h{TIDB_SERVER_IP} -P{TIDB_SERVER_PORT} -u{TIDB_USER} -p{TIDB_USER_PWD} --skip-ssl -e "select * from mysql.user;"'},
    {"cmd_name": "stats_healthy", "cmd_text": f'mysql -h{TIDB_SERVER_IP} -P{TIDB_SERVER_PORT} -u{TIDB_USER} -p{TIDB_USER_PWD} -v -e "SHOW STATS_HEALTHY where Healthy<80;"'},
    {"cmd_name": "show_global_bindings", "cmd_text": f'mysql -h{TIDB_SERVER_IP} -P{TIDB_SERVER_PORT} -u{TIDB_USER} -p{TIDB_USER_PWD} -v -e "show global bindings;"'},
    {"cmd_name": "show_var_tidb_capture_plan_baselines", "cmd_text": f'''mysql -h{TIDB_SERVER_IP} -P{TIDB_SERVER_PORT} -u{TIDB_USER} -p{TIDB_USER_PWD} -e "show variables like 'tidb_capture_plan_baselines';"'''},
    {"cmd_name": "show_var_tidb_evolve_plan_baselines", "cmd_text": f'''mysql -h{TIDB_SERVER_IP} -P{TIDB_SERVER_PORT} -u{TIDB_USER} -p{TIDB_USER_PWD}  -e "show variables like 'tidb_evolve_plan_baselines';"'''},
    {"cmd_name": "tikv_gc_life_time", "cmd_text": f'''mysql -h{TIDB_SERVER_IP} -P{TIDB_SERVER_PORT} -u{TIDB_USER} -p{TIDB_USER_PWD} --skip-column-names -e "select VARIABLE_VALUE from mysql.tidb where VARIABLE_NAME ='tikv_gc_life_time';"'''},
    {"cmd_name": "tikv_gc_run_interval", "cmd_text": f'''mysql -h{TIDB_SERVER_IP} -P{TIDB_SERVER_PORT} -u{TIDB_USER} -p{TIDB_USER_PWD} --skip-column-names -e "select VARIABLE_VALUE from mysql.tidb where VARIABLE_NAME ='tikv_gc_run_interval';"'''},
    {"cmd_name": "tikv_gc_safe_point", "cmd_text": f'''mysql -h{TIDB_SERVER_IP} -P{TIDB_SERVER_PORT} -u{TIDB_USER} -p{TIDB_USER_PWD} --skip-column-names -e "select VARIABLE_VALUE from mysql.tidb where VARIABLE_NAME ='tikv_gc_safe_point';"'''},
    {"cmd_name": "tikv_gc_last_run_time", "cmd_text": f'''mysql -h{TIDB_SERVER_IP} -P{TIDB_SERVER_PORT} -u{TIDB_USER} -p{TIDB_USER_PWD} --skip-column-names -e "select VARIABLE_VALUE from mysql.tidb where VARIABLE_NAME ='tikv_gc_last_run_time';"'''},
    {"cmd_name": "cluster_info", "cmd_text": f'''mysql -h{TIDB_SERVER_IP} -P{TIDB_SERVER_PORT} -u{TIDB_USER} -p{TIDB_USER_PWD} -v --table -e "select type,instance,version,git_hash,uptime from information_schema.cluster_info;"'''},
    {"cmd_name": "binlog_checkpoint", "cmd_text": f'''mysql -h{TIDB_SERVER_IP} -P{TIDB_SERVER_PORT} -u{TIDB_USER} -p{TIDB_USER_PWD} -v -e "select * from tidb_binlog.checkpoint;" || echo ...'''},
    {"cmd_name": "tiflash_replica", "cmd_text": f'''mysql -h{TIDB_SERVER_IP} -P{TIDB_SERVER_PORT} -u{TIDB_USER} -p{TIDB_USER_PWD} -vvv -e --table "select * from information_schema.tiflash_replica;" || echo ...'''}
]

KEYWORDS = [
    {"key_name": "cluster_name", "key_str": f'{CLUSTER_NAME}'}
]

# Download image from Grafana
def download_img(DASHBOARD_ID,PANEL_ID,OUTPUT_FILE):
    headers = {"Authorization": f"Bearer {GRAFANA_API_KEY}"}
    params = {
        "orgId": ORGID,
        "from": FROM_TIME,
        "to": TO_TIME,
        "panelId": PANEL_ID,
        "width": IMG_WIDTH,
        "height": IMG_HEIGHT,
    }
    url = f"{GRAFANA_URL}/render/d-solo/{DASHBOARD_ID}/panel?{requests.compat.urlencode(params)}"
    response = requests.get(url, headers=headers, stream=True)
    if response.status_code == 200:
        with open(f"{OUTPUT_FILE}", "wb") as f:
            for chunk in response.iter_content(1024):
                f.write(chunk)
        #print(f"Image {OUTPUT_FILE} saved successfully")
        logging.info(f"Image {OUTPUT_FILE} saved successfully.")
    else:
        #print(f"Error: {response.status_code}")
        logging.error(f"Error: {response.status_code}")

# Replace placeholder with image in the report
def replace_placeholder_with_image(doc, placeholder, image_path):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, '')
            new_run = paragraph.add_run()
            #new_run.add_picture(image_path, width=width, height=height)
            new_run.add_picture(image_path, width=Cm(25))
            #logging.info(f"Image {image_path} written to document.")
            #for run in paragraph.runs:
            #    #logging.info(f'{run}')
            #    if placeholder in run.text:
            #        run.text = run.text.replace(placeholder, "")
            #        #logging.info(f"Placeholder {placeholder} was removed.")
            #break

# Replace placeholder with query result in report table
def replace_placeholder_in_table(doc, placeholder, cmd_text):
    try:
        result = subprocess.run(cmd_text, capture_output=True, text=True, check=True, shell=True, encoding='utf-8')
        #print(result.stderr)
        #print(result.)
        #if error then

        #
        if placeholder == "{tiup_cluster_check}":
            #  XML compatible
            #command_output_tmp = re.sub(u"[\\x00-\\x08\\x0b\\x0e-\\x1f\\x7f]","",result.stdout)
            #match = re.search('Cleanup check files', result.stdout)
            results = result.stdout + result.stderr
            match = re.search('Node          Check         Result  Message', results)
            if match:
                command_output = re.sub(u"[\\x00-\\x08\\x0b\\x0e-\\x1f\\x7f]","",results[match.start():])
            else:
                command_output = re.sub(u"[\\x00-\\x08\\x0b\\x0e-\\x1f\\x7f]","",results)
        else:
            results = result.stdout
            match = re.search('Outputs', results)
            if match:
                command_output = results[match.start():]
            else:
                command_output = results

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, command_output)
                        logging.info(f"{placeholder} was replaced by result.")
    except subprocess.CalledProcessError as e:
        #print(f"命令执行失败: {e}")
        #command_output = str(e)
        print()
        logging.error(f'Command {placeholder} failed: {e}.')

# Replace placeholder with keywords in report
def replace_placeholder_in_doc(doc, placeholder, key_str):
    try:
        #result = subprocess.run(cmd_text, capture_output=True, text=True, check=True, shell=True, encoding='utf-8')
        #command_output = result.stdout
        #print (command_output)
        for para in doc.paragraphs:
            #for run in paragraph.runs:
            if placeholder in para.text:
                #print(f"{placeholder} replace {key_str}")
                para.text = para.text.replace(placeholder, key_str)
        #logging.info(f"KEYWORDS: {placeholder} was replaced.")
    except subprocess.CalledProcessError as e:
        #print(f"命令执行失败: {e}")
        #command_output = str(e)
        logging.error(f'Command {placeholder} failed: {e}.')

def main():
    # Record start time
    start_time = datetime.datetime.now()
    # Create image folder if not exists
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)
    doc = Document(REPORT_TEMPLATE)
    logging.info("======================================================================")
    logging.info("Start TiDB health check process.")
    print("Start TiDB health check process.")
    # Collecting monitor data
    logging.info("1 - Start replace images in report.")
    print("1 - Start replace images in report.")
    for panel in PANELS_TO_RENDER:
        placeholder = f'{{{panel["output_file"]}}}'
        # print (placeholder)
        output_file_path = os.path.join(OUTPUT_DIR, panel["output_file"])
        download_img(panel["dashboard_id"], panel["panel_id"], output_file_path)
        replace_placeholder_with_image(doc, placeholder, output_file_path)
        logging.info(f"Completed replace images {placeholder} placeholder in report.")
    # Collecting cluster info
    logging.info("2 - Start replace query result in report.")
    print("2 - Start replace query result in report.")
    for cmd in COMMANDS:
        placeholder = f'{{{cmd["cmd_name"]}}}'
        # print(placeholder)
        replace_placeholder_in_table(doc, placeholder, cmd["cmd_text"])
        logging.info(f"Completed replace query result {placeholder} in report.")
    logging.info("3 - Start replace keywords in report.")
    print("3 - Start replace keywords in report.")
    for key in KEYWORDS:
        placeholder = f'{{{key["key_name"]}}}'
        # print(placeholder)
        replace_placeholder_in_doc(doc, placeholder, key["key_str"])
        logging.info(f"Completed replace keywords {placeholder} in report.")
    doc.save(OUTPUT_REPORT)
    # Record end time
    end_time = datetime.datetime.now()
    elapsed_time = end_time - start_time
    logging.info(f"Completed TiDB health check, total elapsed time: {elapsed_time.total_seconds():.2f}s, please check result in report - {OUTPUT_REPORT}.")
    print(f"Completed TiDB health check, total elapsed time: {elapsed_time.total_seconds():.2f}s, please check result in report - {OUTPUT_REPORT}.")
    logging.info("======================================================================")

if __name__ == '__main__':
    main()